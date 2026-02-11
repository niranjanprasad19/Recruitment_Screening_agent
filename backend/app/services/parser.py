"""
RSA MVP Enhanced — File Parser Service
=======================================
Extracts text content from uploaded files (PDF, DOCX, TXT).
Handles various file formats with error recovery.
"""

import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class FileParser:
    """
    Parses uploaded resume/JD files and extracts raw text.
    Supports: PDF, DOCX, TXT, DOC
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".doc"}
    
    @staticmethod
    def parse(file_path: str) -> str:
        """
        Parse a file and return its text content.
        
        Args:
            file_path: Absolute path to the uploaded file.
        
        Returns:
            Extracted text content as a string.
        
        Raises:
            ValueError: If the file format is unsupported.
            FileNotFoundError: If the file doesn't exist.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {FileParser.SUPPORTED_EXTENSIONS}")
        
        try:
            if ext == ".pdf":
                return FileParser._parse_pdf(file_path)
            elif ext == ".docx":
                return FileParser._parse_docx(file_path)
            elif ext == ".txt":
                return FileParser._parse_txt(file_path)
            elif ext == ".doc":
                # Fallback: try DOCX parser (works for some .doc files)
                return FileParser._parse_docx(file_path)
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise ValueError(f"Failed to parse file: {str(e)}")
    
    @staticmethod
    def _parse_pdf(file_path: str) -> str:
        """Extract text from a PDF file using PyPDF2."""
        from PyPDF2 import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception as e:
                logger.warning(f"Could not extract text from page {page_num}: {e}")
                continue
        
        if not text_parts:
            raise ValueError("No text could be extracted from the PDF. It might be image-based.")
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extract text from a DOCX file using python-docx."""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        if not text_parts:
            raise ValueError("No text could be extracted from the DOCX file.")
        
        return "\n".join(text_parts)
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Read text from a plain text file with encoding detection."""
        encodings = ["utf-8", "latin-1", "cp1252", "ascii"]
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                if content.strip():
                    return content
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        raise ValueError("Could not decode the text file with any supported encoding.")
    
    @staticmethod
    def validate_file(filename: str, max_size_bytes: int, file_size: int) -> Optional[str]:
        """
        Validate a file before processing.
        
        Returns:
            None if valid, error message string if invalid.
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in FileParser.SUPPORTED_EXTENSIONS:
            return f"Unsupported file type '{ext}'. Allowed: {', '.join(FileParser.SUPPORTED_EXTENSIONS)}"
        
        if file_size > max_size_bytes:
            max_mb = max_size_bytes / (1024 * 1024)
            return f"File too large ({file_size / (1024*1024):.1f}MB). Maximum: {max_mb}MB"
        
        if file_size == 0:
            return "File is empty."
        
        return None
